# -*- coding: utf-8 -*-
"""
Tests fuer WordReportGenerator.

Testet:
- Report-Erstellung
- Sektionen und Tabellen
- Customer/Contract/Analysis Reports
- Edge Cases und Fehlerbehandlung
"""

import pytest
from datetime import datetime
from pathlib import Path
from io import BytesIO
from unittest.mock import MagicMock, patch
from typing import Dict, Any, List

# Imports mit Fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.services.rag.word_generator import (
    WordReportGenerator,
    get_word_generator,
    DOCX_AVAILABLE as MODULE_DOCX,
)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestWordReportGeneratorInit:
    """Tests fuer Generator-Initialisierung."""

    def test_init_with_docx(self):
        """Sollte Generator initialisieren wenn python-docx verfuegbar."""
        generator = WordReportGenerator()
        assert generator is not None


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestCreateReport:
    """Tests fuer create_report Methode."""

    @pytest.fixture
    def generator(self):
        return WordReportGenerator()

    def test_create_simple_report(self, generator: WordReportGenerator):
        """Sollte einfachen Report erstellen."""
        content = {
            "sections": [
                {
                    "heading": "Testabschnitt",
                    "level": 1,
                    "paragraphs": ["Dies ist ein Testabsatz."]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Test Report", content)

        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify Word content
        doc = Document(BytesIO(result))
        # Titel + Metadaten + Section Heading = mindestens 2 Paragraphen
        assert len(doc.paragraphs) >= 2

    def test_create_report_with_bullet_points(self, generator: WordReportGenerator):
        """Sollte Report mit Aufzaehlungen erstellen."""
        content = {
            "sections": [
                {
                    "heading": "Liste",
                    "bullet_points": ["Punkt 1", "Punkt 2", "Punkt 3"]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Bullet Report", content)
        doc = Document(BytesIO(result))

        # Sollte Bullet-Points enthalten
        assert any("Punkt 1" in p.text for p in doc.paragraphs)

    def test_create_report_with_numbered_list(self, generator: WordReportGenerator):
        """Sollte Report mit nummerierter Liste erstellen."""
        content = {
            "sections": [
                {
                    "heading": "Nummeriert",
                    "numbered_list": ["Erster", "Zweiter", "Dritter"]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Numbered Report", content)
        doc = Document(BytesIO(result))

        assert any("Erster" in p.text for p in doc.paragraphs)

    def test_create_report_with_table(self, generator: WordReportGenerator):
        """Sollte Report mit Tabelle erstellen."""
        content = {
            "sections": [],
            "tables": [
                {
                    "title": "Testdaten",
                    "headers": ["Spalte 1", "Spalte 2", "Spalte 3"],
                    "rows": [
                        ["A", "B", "C"],
                        ["D", "E", "F"],
                    ]
                }
            ]
        }

        result = generator.create_report("Table Report", content)
        doc = Document(BytesIO(result))

        # Sollte mindestens eine Tabelle haben
        assert len(doc.tables) >= 1

        # Tabelle pruefen
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Spalte 1"
        assert table.rows[1].cells[0].text == "A"

    def test_create_report_with_formatted_runs(self, generator: WordReportGenerator):
        """Sollte Report mit formatierten Text-Runs erstellen."""
        content = {
            "sections": [
                {
                    "heading": "Formatiert",
                    "paragraphs": [
                        {
                            "runs": [
                                {"text": "Normal "},
                                {"text": "Fett", "bold": True},
                                {"text": " und "},
                                {"text": "Kursiv", "italic": True}
                            ]
                        }
                    ]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Formatted Report", content)
        doc = Document(BytesIO(result))

        # Sollte formatierten Text enthalten
        found_text = False
        for p in doc.paragraphs:
            if "Fett" in p.text and "Kursiv" in p.text:
                found_text = True
                break
        assert found_text

    def test_create_report_save_to_file(self, generator: WordReportGenerator, tmp_path):
        """Sollte Report in Datei speichern."""
        content = {
            "sections": [{"heading": "Test", "paragraphs": ["Content"]}],
            "tables": []
        }

        output_path = tmp_path / "test_report.docx"
        result = generator.create_report("Test", content, output_path=output_path)

        assert output_path.exists()
        assert output_path.read_bytes() == result

    def test_create_report_with_template(self, generator: WordReportGenerator, tmp_path):
        """Sollte Report mit Vorlage erstellen."""
        # Erstelle einfache Vorlage
        template_doc = Document()
        template_doc.add_heading("Vorlage", 0)
        template_path = tmp_path / "template.docx"
        template_doc.save(template_path)

        content = {
            "sections": [{"heading": "Inhalt", "paragraphs": ["Test"]}],
            "tables": []
        }

        result = generator.create_report(
            "Mit Vorlage",
            content,
            template_path=template_path
        )

        doc = Document(BytesIO(result))
        # Sollte sowohl Vorlage als auch neuen Inhalt haben
        assert any("Vorlage" in p.text for p in doc.paragraphs)
        assert any("Inhalt" in p.text for p in doc.paragraphs)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestCreateCustomerReport:
    """Tests fuer create_customer_report Methode."""

    @pytest.fixture
    def generator(self):
        return WordReportGenerator()

    def test_create_customer_report_minimal(self, generator: WordReportGenerator):
        """Sollte minimalen Kundenreport erstellen."""
        result = generator.create_customer_report(
            customer_name="Test GmbH",
            summary="Zusammenfassung des Kunden.",
            key_facts={},
            documents=[]
        )

        assert isinstance(result, bytes)

        doc = Document(BytesIO(result))
        # Sollte Titel enthalten
        assert any("Test GmbH" in p.text for p in doc.paragraphs)

    def test_create_customer_report_with_data(self, generator: WordReportGenerator):
        """Sollte Kundenreport mit allen Daten erstellen."""
        documents = [
            {"date": "2024-01-15", "type": "Rechnung", "title": "RE-001"},
            {"date": "2024-02-20", "type": "Vertrag", "title": "VT-001"},
        ]
        key_facts = {
            "Kundennummer": "KD-12345",
            "Seit": "2020",
            "Umsatz 2024": "€ 50.000"
        }

        result = generator.create_customer_report(
            customer_name="Muster AG",
            summary="Wichtiger Kunde mit langer Historie.",
            key_facts=key_facts,
            documents=documents
        )

        doc = Document(BytesIO(result))

        # Sollte Zusammenfassung enthalten
        assert any("Wichtiger Kunde" in p.text for p in doc.paragraphs)

        # Sollte Key Facts enthalten
        assert any("Kundennummer" in p.text for p in doc.paragraphs)

        # Sollte Tabelle mit Dokumenten haben
        assert len(doc.tables) >= 1

    def test_create_customer_report_max_20_documents(self, generator: WordReportGenerator):
        """Sollte maximal 20 Dokumente in Tabelle aufnehmen."""
        documents = [
            {"date": f"2024-{i:02d}-01", "type": "Test", "title": f"DOC-{i}"}
            for i in range(1, 30)  # 29 Dokumente
        ]

        result = generator.create_customer_report(
            customer_name="Test",
            summary="Test",
            key_facts={},
            documents=documents
        )

        doc = Document(BytesIO(result))
        table = doc.tables[0]
        # Header + max 20 Dokumente
        assert len(table.rows) <= 21


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestCreateContractReport:
    """Tests fuer create_contract_report Methode."""

    @pytest.fixture
    def generator(self):
        return WordReportGenerator()

    def test_create_contract_report_minimal(self, generator: WordReportGenerator):
        """Sollte minimalen Vertragsreport erstellen."""
        result = generator.create_contract_report(
            contract_title="Rahmenvertrag 2024",
            summary="Vertragszusammenfassung.",
            parties=["Firma A", "Firma B"],
            key_terms=[],
            timeline=[]
        )

        assert isinstance(result, bytes)

        doc = Document(BytesIO(result))
        assert any("Rahmenvertrag" in p.text for p in doc.paragraphs)

    def test_create_contract_report_with_key_terms(self, generator: WordReportGenerator):
        """Sollte Vertragsreport mit Klauseln erstellen."""
        key_terms = [
            {"title": "Laufzeit", "description": "5 Jahre ab Unterzeichnung"},
            {"title": "Kuendigung", "description": "3 Monate zum Quartalsende"},
        ]

        result = generator.create_contract_report(
            contract_title="Servicevertrag",
            summary="IT-Servicevertrag.",
            parties=["Kunde GmbH", "Dienstleister AG"],
            key_terms=key_terms,
            timeline=[]
        )

        doc = Document(BytesIO(result))
        # Sollte Klausel-Titel enthalten
        assert any("Laufzeit" in p.text for p in doc.paragraphs)

    def test_create_contract_report_with_timeline(self, generator: WordReportGenerator):
        """Sollte Vertragsreport mit Timeline erstellen."""
        timeline = [
            {"date": "2024-01-01", "event": "Vertragsbeginn", "description": "Start"},
            {"date": "2024-12-31", "event": "Vertragsende", "description": "Ende der Laufzeit"},
        ]

        result = generator.create_contract_report(
            contract_title="Projektvertrag",
            summary="Projektzusammenfassung.",
            parties=["A", "B"],
            key_terms=[],
            timeline=timeline
        )

        doc = Document(BytesIO(result))
        # Sollte Timeline-Tabelle haben
        assert len(doc.tables) >= 1


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestCreateAnalysisReport:
    """Tests fuer create_analysis_report Methode."""

    @pytest.fixture
    def generator(self):
        return WordReportGenerator()

    def test_create_analysis_report_minimal(self, generator: WordReportGenerator):
        """Sollte minimalen Analyse-Report erstellen."""
        result = generator.create_analysis_report(
            title="Marktanalyse Q4 2024",
            executive_summary="Positive Marktentwicklung erwartet.",
            findings=[],
            recommendations=[]
        )

        assert isinstance(result, bytes)

        doc = Document(BytesIO(result))
        assert any("Marktanalyse" in p.text for p in doc.paragraphs)

    def test_create_analysis_report_with_findings(self, generator: WordReportGenerator):
        """Sollte Analyse-Report mit Ergebnissen erstellen."""
        findings = [
            {
                "title": "Wachstum",
                "description": "Umsatzwachstum von 15% im Vergleich zum Vorjahr.",
                "details": ["Region Nord: +20%", "Region Sued: +10%"]
            },
            {
                "title": "Marktanteil",
                "description": "Marktanteil auf 25% gestiegen.",
                "details": []
            }
        ]

        result = generator.create_analysis_report(
            title="Jahresanalyse",
            executive_summary="Erfolgreiches Jahr.",
            findings=findings,
            recommendations=["Investitionen fortsetzen", "Neue Maerkte erschliessen"]
        )

        doc = Document(BytesIO(result))
        # Sollte Findings enthalten
        assert any("Wachstum" in p.text for p in doc.paragraphs)
        # Sollte Empfehlungen enthalten
        assert any("Investitionen" in p.text for p in doc.paragraphs)

    def test_create_analysis_report_with_appendix(self, generator: WordReportGenerator):
        """Sollte Analyse-Report mit Anhang erstellen."""
        appendix = {
            "paragraphs": ["Zusaetzliche Informationen hier."],
            "tables": [
                {
                    "title": "Rohdaten",
                    "headers": ["Monat", "Umsatz"],
                    "rows": [["Januar", "100k"], ["Februar", "120k"]]
                }
            ]
        }

        result = generator.create_analysis_report(
            title="Report mit Anhang",
            executive_summary="Summary.",
            findings=[],
            recommendations=[],
            appendix=appendix
        )

        doc = Document(BytesIO(result))
        # Sollte Anhang-Sektion haben
        assert any("Anhang" in p.text for p in doc.paragraphs)
        # Sollte Tabelle haben
        assert len(doc.tables) >= 1


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx nicht installiert")
class TestEdgeCases:
    """Tests fuer Randfaelle."""

    @pytest.fixture
    def generator(self):
        return WordReportGenerator()

    def test_empty_content(self, generator: WordReportGenerator):
        """Sollte leeren Inhalt verarbeiten."""
        content = {
            "sections": [],
            "tables": []
        }

        result = generator.create_report("Empty Report", content)
        assert isinstance(result, bytes)

    def test_empty_table(self, generator: WordReportGenerator):
        """Sollte leere Tabelle ignorieren."""
        content = {
            "sections": [],
            "tables": [
                {"headers": [], "rows": []}
            ]
        }

        result = generator.create_report("Empty Table", content)
        doc = Document(BytesIO(result))
        # Keine Tabelle sollte erstellt werden
        assert len(doc.tables) == 0

    def test_table_with_none_values(self, generator: WordReportGenerator):
        """Sollte None-Werte in Tabellen behandeln."""
        content = {
            "sections": [],
            "tables": [
                {
                    "headers": ["A", "B"],
                    "rows": [[None, "Value"], ["Test", None]]
                }
            ]
        }

        result = generator.create_report("Null Values", content)
        doc = Document(BytesIO(result))
        assert len(doc.tables) == 1

    def test_unicode_content(self, generator: WordReportGenerator):
        """Sollte Unicode-Inhalte verarbeiten."""
        content = {
            "sections": [
                {
                    "heading": "Übersicht mit Ümläuten",
                    "paragraphs": [
                        "Größe, Müller, Böhm, Fälle",
                        "日本語テスト"  # Japanisch
                    ]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Unicode Report", content)
        doc = Document(BytesIO(result))
        # Sollte Umlaute enthalten
        assert any("Größe" in p.text for p in doc.paragraphs)

    def test_very_long_paragraph(self, generator: WordReportGenerator):
        """Sollte sehr lange Absaetze verarbeiten."""
        long_text = "Lorem ipsum " * 1000

        content = {
            "sections": [
                {
                    "heading": "Lang",
                    "paragraphs": [long_text]
                }
            ],
            "tables": []
        }

        result = generator.create_report("Long Report", content)
        assert isinstance(result, bytes)

    def test_multiple_heading_levels(self, generator: WordReportGenerator):
        """Sollte verschiedene Ueberschriften-Level verarbeiten."""
        content = {
            "sections": [
                {"heading": "Level 1", "level": 1, "paragraphs": []},
                {"heading": "Level 2", "level": 2, "paragraphs": []},
                {"heading": "Level 3", "level": 3, "paragraphs": []},
            ],
            "tables": []
        }

        result = generator.create_report("Multi-Level", content)
        doc = Document(BytesIO(result))

        # Alle Levels sollten vorhanden sein
        texts = [p.text for p in doc.paragraphs]
        assert "Level 1" in texts
        assert "Level 2" in texts
        assert "Level 3" in texts


class TestSingleton:
    """Tests fuer Singleton-Pattern."""

    def test_get_word_generator_singleton(self):
        """Sollte immer gleiche Instanz zurueckgeben."""
        # Reset singleton
        import app.services.rag.word_generator as module
        module._word_generator = None

        gen1 = get_word_generator()
        gen2 = get_word_generator()

        assert gen1 is gen2


class TestWithoutDocx:
    """Tests wenn python-docx nicht verfuegbar ist."""

    def test_create_report_raises_without_docx(self):
        """Sollte ImportError werfen wenn python-docx fehlt."""
        with patch.object(
            WordReportGenerator, '__init__',
            return_value=None
        ):
            generator = WordReportGenerator.__new__(WordReportGenerator)

            with patch(
                'app.services.rag.word_generator.DOCX_AVAILABLE',
                False
            ):
                with pytest.raises(ImportError, match="python-docx"):
                    generator.create_report("Test", {})
