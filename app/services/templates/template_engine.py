"""
Document Template Engine Service.

Generiert PDF/DOCX/HTML-Dokumente aus Jinja2-Templates mit deutschen Vorlagen.

Features:
- Jinja2-basierte Template-Engine
- Multi-Format-Export (PDF, DOCX, HTML)
- Built-in deutsche Templates (Rechnung, Angebot, Mahnung, etc.)
- Variable-Validierung und Default-Werte
- WeasyPrint für PDF-Generierung
- python-docx für DOCX-Generierung
"""

import io
import json
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Dict, List, Optional, Union
from uuid import UUID

import structlog
from jinja2 import (
    Environment,
    FileSystemLoader,
    TemplateNotFound,
    select_autoescape,
)
from sqlalchemy.ext.asyncio import AsyncSession

# Optional imports for PDF/DOCX generation
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    HTML = None  # type: ignore
    WEASYPRINT_AVAILABLE = False
    structlog.get_logger(__name__).warning("weasyprint not installed - PDF generation disabled")

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None  # type: ignore
    Inches = Pt = WD_PARAGRAPH_ALIGNMENT = None  # type: ignore
    DOCX_AVAILABLE = False
    structlog.get_logger(__name__).warning("python-docx not installed - DOCX generation disabled")

from app.core.security.sensitive_data_filter import get_pii_safe_logger

logger = get_pii_safe_logger(__name__)

# Template categories
TEMPLATE_CATEGORIES = [
    "rechnung",
    "angebot",
    "mahnung",
    "gutschrift",
    "lieferschein",
    "vertrag",
    "brief",
]


# ============================================================================
# DATA CLASSES
# ============================================================================


@dataclass
class RenderedDocument:
    """Gerendertes Dokument."""

    content: bytes
    filename: str
    mime_type: str
    format: str  # pdf, docx, html
    template_id: str


@dataclass
class TemplateInfo:
    """Template-Informationen."""

    id: str
    name: str  # German
    category: str
    description: str  # German
    variables: List[str]
    formats: List[str]


@dataclass
class TemplateVariable:
    """Template-Variable Definition."""

    name: str
    label: str  # German
    type: str  # text, number, date, currency, address
    required: bool
    default: Optional[str] = None


# ============================================================================
# BUILT-IN TEMPLATES REGISTRY
# ============================================================================

BUILT_IN_TEMPLATES = {
    "rechnung_standard": TemplateInfo(
        id="rechnung_standard",
        name="Standard-Rechnung",
        category="rechnung",
        description="Standardvorlage für Ausgangsrechnungen mit deutscher Formatierung",
        variables=[
            "firma_name",
            "firma_strasse",
            "firma_plz",
            "firma_ort",
            "firma_email",
            "firma_telefon",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "rechnungsnummer",
            "rechnungsdatum",
            "leistungsdatum",
            "zahlungsziel",
            "positionen",  # JSON array
            "nettobetrag",
            "umsatzsteuer",
            "bruttobetrag",
            "bankverbindung",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "angebot_standard": TemplateInfo(
        id="angebot_standard",
        name="Standard-Angebot",
        category="angebot",
        description="Standardvorlage für Angebote",
        variables=[
            "firma_name",
            "firma_strasse",
            "firma_plz",
            "firma_ort",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "angebotsnummer",
            "angebotsdatum",
            "gueltig_bis",
            "positionen",
            "nettobetrag",
            "umsatzsteuer",
            "bruttobetrag",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "mahnung_1": TemplateInfo(
        id="mahnung_1",
        name="1. Mahnung",
        category="mahnung",
        description="Erste Mahnung (freundliche Zahlungserinnerung)",
        variables=[
            "firma_name",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "rechnungsnummer",
            "rechnungsdatum",
            "faelligkeit",
            "betrag",
            "mahngebuehr",
            "gesamtbetrag",
            "bankverbindung",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "mahnung_2": TemplateInfo(
        id="mahnung_2",
        name="2. Mahnung",
        category="mahnung",
        description="Zweite Mahnung (dringliche Zahlungsaufforderung)",
        variables=[
            "firma_name",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "rechnungsnummer",
            "rechnungsdatum",
            "faelligkeit",
            "betrag",
            "mahngebuehr",
            "verzugszinsen",
            "gesamtbetrag",
            "bankverbindung",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "mahnung_3": TemplateInfo(
        id="mahnung_3",
        name="3. Mahnung (Letzte Mahnung)",
        category="mahnung",
        description="Dritte Mahnung (letzte Mahnung vor rechtlichen Schritten)",
        variables=[
            "firma_name",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "rechnungsnummer",
            "rechnungsdatum",
            "faelligkeit",
            "betrag",
            "mahngebuehr",
            "verzugszinsen",
            "inkassokosten",
            "gesamtbetrag",
            "frist_tage",
            "bankverbindung",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "gutschrift_standard": TemplateInfo(
        id="gutschrift_standard",
        name="Standard-Gutschrift",
        category="gutschrift",
        description="Standardvorlage für Gutschriften",
        variables=[
            "firma_name",
            "firma_strasse",
            "firma_plz",
            "firma_ort",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "gutschriftsnummer",
            "gutschriftsdatum",
            "ursprungsrechnung",
            "positionen",
            "nettobetrag",
            "umsatzsteuer",
            "bruttobetrag",
        ],
        formats=["pdf", "docx", "html"],
    ),
    "lieferschein_standard": TemplateInfo(
        id="lieferschein_standard",
        name="Standard-Lieferschein",
        category="lieferschein",
        description="Standardvorlage für Lieferscheine",
        variables=[
            "firma_name",
            "kunde_firma",
            "kunde_strasse",
            "kunde_plz",
            "kunde_ort",
            "lieferscheinnummer",
            "lieferdatum",
            "bestellnummer",
            "positionen",
        ],
        formats=["pdf", "docx", "html"],
    ),
}


# ============================================================================
# TEMPLATE ENGINE SERVICE
# ============================================================================


class TemplateEngineService:
    """Service für Template-basierte Dokumentgenerierung."""

    def __init__(self) -> None:
        """Initialisiert Template-Engine mit Jinja2."""
        # Templates-Verzeichnis (erstelle bei Bedarf)
        self.templates_dir = Path(__file__).parent / "templates"
        self.templates_dir.mkdir(exist_ok=True)

        # Jinja2-Environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(["html", "xml"]),
        )

        # Custom filters
        self.jinja_env.filters["currency"] = self._format_currency
        self.jinja_env.filters["date_de"] = self._format_date_de
        self.jinja_env.filters["number_de"] = self._format_number_de

    # ========================================================================
    # CUSTOM FILTERS
    # ========================================================================

    @staticmethod
    def _format_currency(value: object) -> str:
        """Formatiert Währung in deutschem Format."""
        if value is None:
            return "0,00 €"
        try:
            decimal_value = Decimal(str(value))
            formatted = f"{decimal_value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            return f"{formatted} €"
        except (ValueError, TypeError):
            return "0,00 €"

    @staticmethod
    def _format_date_de(value: object) -> str:
        """Formatiert Datum in deutschem Format."""
        if value is None:
            return ""
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace("Z", "+00:00"))
            except ValueError:
                return value
        if isinstance(value, datetime):
            return value.strftime("%d.%m.%Y")
        return str(value)

    @staticmethod
    def _format_number_de(value: object) -> str:
        """Formatiert Zahl in deutschem Format."""
        if value is None:
            return "0"
        try:
            decimal_value = Decimal(str(value))
            formatted = f"{decimal_value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            return formatted
        except (ValueError, TypeError):
            return str(value)

    # ========================================================================
    # PUBLIC API
    # ========================================================================

    async def render_template(
        self,
        template_id: str,
        data: Dict[str, object],
        output_format: str,
        db: AsyncSession,
    ) -> RenderedDocument:
        """
        Rendert Template mit Daten.

        Args:
            template_id: Template-ID (z.B. "rechnung_standard")
            data: Template-Daten (Variablen)
            output_format: Ausgabeformat (pdf, docx, html)
            db: Datenbank-Session

        Returns:
            RenderedDocument mit Content und Metadaten

        Raises:
            ValueError: Template nicht gefunden oder ungültiges Format
        """
        logger.info(
            "template_render_started",
            template_id=template_id,
            output_format=output_format,
        )

        # Validierung
        if template_id not in BUILT_IN_TEMPLATES:
            raise ValueError(f"Template nicht gefunden: {template_id}")

        template_info = BUILT_IN_TEMPLATES[template_id]
        if output_format not in template_info.formats:
            raise ValueError(
                f"Format '{output_format}' nicht unterstützt für Template '{template_id}'"
            )

        # Variablen validieren
        self._validate_variables(template_info, data)

        # HTML rendern
        html_content = await self._render_html(template_id, data)

        # Format-spezifische Konvertierung
        if output_format == "html":
            content = html_content.encode("utf-8")
            mime_type = "text/html"
        elif output_format == "pdf":
            content = await self._convert_to_pdf(html_content)
            mime_type = "application/pdf"
        elif output_format == "docx":
            content = await self._convert_to_docx(template_id, data)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise ValueError(f"Unbekanntes Format: {output_format}")

        # Filename generieren
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{template_id}_{timestamp}.{output_format}"

        logger.info(
            "template_rendered",
            template_id=template_id,
            format=output_format,
            size_bytes=len(content),
        )

        return RenderedDocument(
            content=content,
            filename=filename,
            mime_type=mime_type,
            format=output_format,
            template_id=template_id,
        )

    async def list_templates(
        self, category: Optional[str] = None, db: AsyncSession = None
    ) -> List[TemplateInfo]:
        """
        Listet verfügbare Templates auf.

        Args:
            category: Optional filter nach Kategorie
            db: Datenbank-Session (für spätere Custom-Templates)

        Returns:
            Liste von TemplateInfo-Objekten
        """
        templates = list(BUILT_IN_TEMPLATES.values())

        if category:
            templates = [t for t in templates if t.category == category]

        logger.debug("templates_listed", count=len(templates), category=category)
        return templates

    async def get_template_variables(self, template_id: str) -> List[TemplateVariable]:
        """
        Gibt Template-Variablen mit Metadaten zurück.

        Args:
            template_id: Template-ID

        Returns:
            Liste von TemplateVariable-Definitionen

        Raises:
            ValueError: Template nicht gefunden
        """
        if template_id not in BUILT_IN_TEMPLATES:
            raise ValueError(f"Template nicht gefunden: {template_id}")

        # Variable-Definitionen basierend auf Template
        variable_defs = self._get_variable_definitions(template_id)

        logger.debug(
            "template_variables_retrieved",
            template_id=template_id,
            count=len(variable_defs),
        )
        return variable_defs

    # ========================================================================
    # PRIVATE HELPERS
    # ========================================================================

    def _validate_variables(
        self, template_info: TemplateInfo, data: Dict[str, object]
    ) -> None:
        """Validiert Template-Variablen."""
        variable_defs = self._get_variable_definitions(template_info.id)
        required_vars = [v.name for v in variable_defs if v.required]

        missing = [var for var in required_vars if var not in data]
        if missing:
            raise ValueError(
                f"Fehlende erforderliche Variablen: {', '.join(missing)}"
            )

    async def _render_html(self, template_id: str, data: Dict[str, object]) -> str:
        """Rendert HTML mit Jinja2."""
        # Default-Template wenn keine Custom-Template-Datei vorhanden
        html_template = self._get_default_html_template(template_id)

        template = self.jinja_env.from_string(html_template)
        return template.render(**data)

    async def _convert_to_pdf(self, html_content: str) -> bytes:
        """Konvertiert HTML zu PDF mit WeasyPrint."""
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

    async def _convert_to_docx(
        self, template_id: str, data: Dict[str, object]
    ) -> bytes:
        """Konvertiert zu DOCX mit python-docx."""
        doc = DocxDocument()

        # Header mit Firmenname
        if "firma_name" in data:
            heading = doc.add_heading(data["firma_name"], 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Template-spezifischer Content
        if template_id.startswith("rechnung"):
            self._add_rechnung_content(doc, data)
        elif template_id.startswith("mahnung"):
            self._add_mahnung_content(doc, data)
        # ... weitere Template-Typen

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_rechnung_content(
        self, doc: DocxDocument, data: Dict[str, object]
    ) -> None:
        """Fügt Rechnungs-Content hinzu."""
        # Empfänger
        doc.add_paragraph(f"{data.get('kunde_firma', '')}")
        doc.add_paragraph(f"{data.get('kunde_strasse', '')}")
        doc.add_paragraph(
            f"{data.get('kunde_plz', '')} {data.get('kunde_ort', '')}"
        )
        doc.add_paragraph("")

        # Rechnungsdaten
        doc.add_heading("Rechnung", 1)
        doc.add_paragraph(f"Rechnungsnummer: {data.get('rechnungsnummer', '')}")
        doc.add_paragraph(f"Rechnungsdatum: {data.get('rechnungsdatum', '')}")
        doc.add_paragraph(f"Leistungsdatum: {data.get('leistungsdatum', '')}")
        doc.add_paragraph("")

        # Positionen (einfache Tabelle)
        positionen = data.get("positionen", [])
        if positionen:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Pos."
            hdr_cells[1].text = "Bezeichnung"
            hdr_cells[2].text = "Menge"
            hdr_cells[3].text = "Betrag"

            for i, pos in enumerate(positionen, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = pos.get("bezeichnung", "")
                row_cells[2].text = str(pos.get("menge", ""))
                row_cells[3].text = self._format_currency(pos.get("betrag", 0))

        doc.add_paragraph("")

        # Summen
        doc.add_paragraph(
            f"Nettobetrag: {self._format_currency(data.get('nettobetrag', 0))}"
        )
        doc.add_paragraph(
            f"Umsatzsteuer: {self._format_currency(data.get('umsatzsteuer', 0))}"
        )
        doc.add_paragraph(
            f"Bruttobetrag: {self._format_currency(data.get('bruttobetrag', 0))}"
        ).bold = True

    def _add_mahnung_content(
        self, doc: DocxDocument, data: Dict[str, object]
    ) -> None:
        """Fügt Mahnungs-Content hinzu."""
        # Empfänger
        doc.add_paragraph(f"{data.get('kunde_firma', '')}")
        doc.add_paragraph(f"{data.get('kunde_strasse', '')}")
        doc.add_paragraph(
            f"{data.get('kunde_plz', '')} {data.get('kunde_ort', '')}"
        )
        doc.add_paragraph("")

        # Mahnung
        doc.add_heading("Zahlungserinnerung", 1)
        doc.add_paragraph(
            f"Sehr geehrte Damen und Herren,\n\n"
            f"zu unserer Rechnung Nr. {data.get('rechnungsnummer', '')} vom "
            f"{data.get('rechnungsdatum', '')} konnten wir bisher keinen Zahlungseingang "
            f"feststellen.\n\n"
            f"Fälligkeit: {data.get('faelligkeit', '')}\n"
            f"Offener Betrag: {self._format_currency(data.get('betrag', 0))}\n\n"
            f"Bitte überweisen Sie den Betrag umgehend auf unser Konto.\n\n"
            f"Mit freundlichen Grüßen"
        )

    def _get_variable_definitions(self, template_id: str) -> List[TemplateVariable]:
        """Gibt Variable-Definitionen für Template zurück."""
        # Gemeinsame Variablen
        common_vars = [
            TemplateVariable(
                name="firma_name",
                label="Firmenname",
                type="text",
                required=True,
            ),
            TemplateVariable(
                name="firma_strasse",
                label="Straße (Firma)",
                type="text",
                required=False,
            ),
            TemplateVariable(
                name="firma_plz",
                label="PLZ (Firma)",
                type="text",
                required=False,
            ),
            TemplateVariable(
                name="firma_ort",
                label="Ort (Firma)",
                type="text",
                required=False,
            ),
            TemplateVariable(
                name="kunde_firma",
                label="Kundenname",
                type="text",
                required=True,
            ),
            TemplateVariable(
                name="kunde_strasse",
                label="Straße (Kunde)",
                type="text",
                required=False,
            ),
            TemplateVariable(
                name="kunde_plz",
                label="PLZ (Kunde)",
                type="text",
                required=False,
            ),
            TemplateVariable(
                name="kunde_ort",
                label="Ort (Kunde)",
                type="text",
                required=False,
            ),
        ]

        # Template-spezifische Variablen
        if template_id == "rechnung_standard":
            return common_vars + [
                TemplateVariable(
                    name="rechnungsnummer",
                    label="Rechnungsnummer",
                    type="text",
                    required=True,
                ),
                TemplateVariable(
                    name="rechnungsdatum",
                    label="Rechnungsdatum",
                    type="date",
                    required=True,
                ),
                TemplateVariable(
                    name="nettobetrag",
                    label="Nettobetrag",
                    type="currency",
                    required=True,
                ),
                TemplateVariable(
                    name="bruttobetrag",
                    label="Bruttobetrag",
                    type="currency",
                    required=True,
                ),
            ]

        # Fallback
        return common_vars

    def _get_default_html_template(self, template_id: str) -> str:
        """Gibt Default-HTML-Template zurück."""
        # Einfaches Basis-Template
        base_html = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                .header { text-align: center; margin-bottom: 30px; }
                .address { margin: 20px 0; }
                .content { margin: 20px 0; }
                table { width: 100%; border-collapse: collapse; margin: 20px 0; }
                th, td { padding: 8px; text-align: left; border-bottom: 1px solid #ddd; }
                .total { font-weight: bold; font-size: 1.2em; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{{ firma_name }}</h1>
                <p>{{ firma_strasse }}, {{ firma_plz }} {{ firma_ort }}</p>
            </div>

            <div class="address">
                <p><strong>{{ kunde_firma }}</strong></p>
                <p>{{ kunde_strasse }}</p>
                <p>{{ kunde_plz }} {{ kunde_ort }}</p>
            </div>

            <div class="content">
                <!-- Template-spezifischer Content -->
                {% block content %}{% endblock %}
            </div>
        </body>
        </html>
        """
        return base_html
