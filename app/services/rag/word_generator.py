"""Word Report Generator.

Generiert Word-Dokumente mit:
- Strukturierte Abschnitte
- Tabellen
- Formatierung
- LLM-gestuetzte Inhalte
"""

import structlog
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = structlog.get_logger(__name__)


class WordReportGenerator:
    """Generator für Word-Reports."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            logger.warning("python-docx_not_available")

    def create_report(
        self,
        title: str,
        content: Dict[str, Any],
        template_path: Optional[Path] = None,
        output_path: Optional[Path] = None
    ) -> bytes:
        """
        Erstellt ein Word-Dokument.

        Args:
            title: Dokumenttitel
            content: Dokumentinhalt mit Sektionen
            template_path: Optionale Vorlage
            output_path: Optionaler Speicherpfad

        Returns:
            Word-Datei als Bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx ist nicht installiert. "
                "Installieren mit: pip install python-docx"
            )

        # Dokument erstellen (mit oder ohne Vorlage)
        if template_path and template_path.exists():
            doc = Document(template_path)
        else:
            doc = Document()
            self._setup_styles(doc)

        # Titel
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadaten
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Erstellt am {datetime.now().strftime('%d.%m.%Y um %H:%M Uhr')}"
        )
        meta_run.italic = True
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Abstand

        # Sektionen hinzufuegen
        for section in content.get("sections", []):
            self._add_section(doc, section)

        # Tabellen hinzufuegen
        for table_config in content.get("tables", []):
            self._add_table(doc, table_config)

        # Als Bytes speichern
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        file_content = output.read()

        # Optional auf Disk speichern
        if output_path:
            output_path.write_bytes(file_content)
            logger.info("word_report_saved", path=str(output_path))

        return file_content

    def _setup_styles(self, doc: "Document"):
        """Richtet Standard-Styles ein."""
        styles = doc.styles

        # Heading 1 Style
        if "Heading 1" in [s.name for s in styles]:
            h1_style = styles["Heading 1"]
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(0, 51, 102)

        # Heading 2 Style
        if "Heading 2" in [s.name for s in styles]:
            h2_style = styles["Heading 2"]
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(0, 102, 153)

    def _add_section(self, doc: "Document", section: Dict[str, Any]):
        """Fuegt eine Sektion hinzu."""
        # Überschrift
        if section.get("heading"):
            level = section.get("level", 1)
            doc.add_heading(section["heading"], level=level)

        # Text-Absätze
        for paragraph in section.get("paragraphs", []):
            if isinstance(paragraph, str):
                doc.add_paragraph(paragraph)
            elif isinstance(paragraph, dict):
                p = doc.add_paragraph()
                for run_config in paragraph.get("runs", []):
                    run = p.add_run(run_config.get("text", ""))
                    if run_config.get("bold"):
                        run.bold = True
                    if run_config.get("italic"):
                        run.italic = True
                    if run_config.get("underline"):
                        run.underline = True

        # Aufzaehlung
        for item in section.get("bullet_points", []):
            doc.add_paragraph(item, style="List Bullet")

        # Nummerierte Liste
        for item in section.get("numbered_list", []):
            doc.add_paragraph(item, style="List Number")

        doc.add_paragraph()  # Abstand

    def _add_table(self, doc: "Document", config: Dict[str, Any]):
        """Fuegt eine Tabelle hinzu."""
        headers = config.get("headers", [])
        rows = config.get("rows", [])
        title = config.get("title")

        if not headers or not rows:
            return

        # Tabellen-Titel
        if title:
            doc.add_heading(title, level=2)

        # Tabelle erstellen
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            # Header fett machen
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Daten
        for row_data in rows:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row_data):
                if idx < len(row_cells):
                    row_cells[idx].text = str(value) if value is not None else ""

        doc.add_paragraph()  # Abstand

    def create_customer_report(
        self,
        customer_name: str,
        summary: str,
        key_facts: Dict[str, Any],
        documents: List[Dict[str, Any]],
        output_path: Optional[Path] = None
    ) -> bytes:
        """
        Erstellt einen Kunden-Report als Word-Dokument.

        Args:
            customer_name: Kundenname
            summary: LLM-generierte Zusammenfassung
            key_facts: Wichtige Kennzahlen
            documents: Dokumentenliste
            output_path: Optionaler Speicherpfad

        Returns:
            Word-Datei als Bytes
        """
        content = {
            "sections": [
                {
                    "heading": "Zusammenfassung",
                    "level": 1,
                    "paragraphs": [summary]
                },
                {
                    "heading": "Wichtige Fakten",
                    "level": 1,
                    "bullet_points": [
                        f"{k}: {v}" for k, v in key_facts.items()
                    ]
                },
                {
                    "heading": "Dokumentenübersicht",
                    "level": 1,
                    "paragraphs": [
                        f"Insgesamt {len(documents)} Dokumente im System."
                    ]
                }
            ],
            "tables": [
                {
                    "title": "Dokumentenliste",
                    "headers": ["Datum", "Typ", "Titel"],
                    "rows": [
                        [
                            d.get("date", "-"),
                            d.get("type", "-"),
                            d.get("title", "-")
                        ]
                        for d in documents[:20]  # Max 20 Dokumente
                    ]
                }
            ]
        }

        return self.create_report(
            title=f"Kundenreport: {customer_name}",
            content=content,
            output_path=output_path
        )

    def create_contract_report(
        self,
        contract_title: str,
        summary: str,
        parties: List[str],
        key_terms: List[Dict[str, Any]],
        timeline: List[Dict[str, Any]],
        output_path: Optional[Path] = None
    ) -> bytes:
        """
        Erstellt einen Vertrags-Report.

        Args:
            contract_title: Vertragstitel
            summary: LLM-generierte Zusammenfassung
            parties: Vertragsparteien
            key_terms: Wichtige Vertragsklauseln
            timeline: Wichtige Termine
            output_path: Optionaler Speicherpfad

        Returns:
            Word-Datei als Bytes
        """
        content = {
            "sections": [
                {
                    "heading": "Vertragszusammenfassung",
                    "level": 1,
                    "paragraphs": [summary]
                },
                {
                    "heading": "Vertragsparteien",
                    "level": 1,
                    "numbered_list": parties
                },
                {
                    "heading": "Wichtige Vertragsklauseln",
                    "level": 1,
                    "paragraphs": []
                }
            ],
            "tables": []
        }

        # Key Terms als Unterabschnitte
        for term in key_terms:
            content["sections"].append({
                "heading": term.get("title", "Klausel"),
                "level": 2,
                "paragraphs": [term.get("description", "")]
            })

        # Timeline Tabelle
        if timeline:
            content["tables"].append({
                "title": "Wichtige Termine",
                "headers": ["Datum", "Ereignis", "Beschreibung"],
                "rows": [
                    [
                        t.get("date", "-"),
                        t.get("event", "-"),
                        t.get("description", "-")
                    ]
                    for t in timeline
                ]
            })

        return self.create_report(
            title=f"Vertragsreport: {contract_title}",
            content=content,
            output_path=output_path
        )

    def create_analysis_report(
        self,
        title: str,
        executive_summary: str,
        findings: List[Dict[str, Any]],
        recommendations: List[str],
        appendix: Optional[Dict[str, Any]] = None,
        output_path: Optional[Path] = None
    ) -> bytes:
        """
        Erstellt einen Analyse-Report.

        Args:
            title: Report-Titel
            executive_summary: Management Summary
            findings: Analyseergebnisse
            recommendations: Empfehlungen
            appendix: Optionaler Anhang
            output_path: Optionaler Speicherpfad

        Returns:
            Word-Datei als Bytes
        """
        content = {
            "sections": [
                {
                    "heading": "Management Summary",
                    "level": 1,
                    "paragraphs": [executive_summary]
                },
                {
                    "heading": "Analyseergebnisse",
                    "level": 1,
                    "paragraphs": []
                }
            ],
            "tables": []
        }

        # Findings als Unterabschnitte
        for idx, finding in enumerate(findings, 1):
            content["sections"].append({
                "heading": f"Ergebnis {idx}: {finding.get('title', '')}",
                "level": 2,
                "paragraphs": [finding.get("description", "")],
                "bullet_points": finding.get("details", [])
            })

        # Empfehlungen
        content["sections"].append({
            "heading": "Empfehlungen",
            "level": 1,
            "numbered_list": recommendations
        })

        # Anhang
        if appendix:
            content["sections"].append({
                "heading": "Anhang",
                "level": 1,
                "paragraphs": appendix.get("paragraphs", [])
            })
            if appendix.get("tables"):
                content["tables"].extend(appendix["tables"])

        return self.create_report(
            title=title,
            content=content,
            output_path=output_path
        )


# Singleton
_word_generator: Optional[WordReportGenerator] = None


def get_word_generator() -> WordReportGenerator:
    """Gibt WordReportGenerator Singleton zurück."""
    global _word_generator
    if _word_generator is None:
        _word_generator = WordReportGenerator()
    return _word_generator
